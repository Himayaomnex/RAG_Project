"""
================================================================================
Pipeline Module (pipeline.py)
================================================================================
Defines the shared vector database, semantic chunking parser, and the exact 4
retrieval pipelines developed by the team:

  P1 — Scroll + Custom Reranker (Document-balanced top 2/doc, top 15 + Heuristic Reranker)
  P2 — Scroll API Scan (Document-balanced top 4/doc, top 35 candidates, no reranker)
  P3 — Top-K Vector Search (Single-shot Qdrant query_points, fixed K=15)
  P4 — Two-Stage Map-Reduce (100% full-corpus scroll grouped by file + Map summarization + Reduce synthesis)
================================================================================
"""

import os
import re
import uuid
import shelve
import hashlib
import json
import docx
from typing import List, Dict, Any, Optional
import torch
from dotenv import load_dotenv
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams, PointStruct, Filter, FieldCondition, MatchText, MatchValue
from sentence_transformers import SentenceTransformer, util

load_dotenv()


class CachedEmbeddingModel:
    """Wraps SentenceTransformer with a local disk cache to prevent slow re-embedding."""
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.model = SentenceTransformer(model_name)
        self.cache_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "emb_cache")
        
    def encode(self, texts, convert_to_tensor: bool = False):
        is_single = isinstance(texts, str)
        text_list = [texts] if is_single else texts
        
        embeddings = []
        needs_encoding = []
        needs_encoding_idx = []
        
        try:
            with shelve.open(self.cache_file) as cache:
                for i, t in enumerate(text_list):
                    h = hashlib.md5(t.encode('utf-8')).hexdigest()
                    if h in cache:
                        embeddings.append(cache[h])
                    else:
                        embeddings.append(None)
                        needs_encoding.append(t)
                        needs_encoding_idx.append(i)
                        
                if needs_encoding:
                    fresh_embs = self.model.encode(needs_encoding).tolist()
                    for idx, emb in zip(needs_encoding_idx, fresh_embs):
                        embeddings[idx] = emb
                        cache[hashlib.md5(text_list[idx].encode('utf-8')).hexdigest()] = emb
        except Exception:
            # Fallback if cache is locked or unavailable
            embeddings = self.model.encode(text_list).tolist()
                    
        if convert_to_tensor:
            return torch.tensor(embeddings)
        return embeddings[0] if is_single else embeddings


class Sentence:
    def __init__(self, text: str, speaker: str, page: int, file_path: str):
        self.text = text
        self.speaker = speaker
        self.page = page
        self.file_path = file_path


class SemanticTranscriptParser:
    """Parses .docx transcripts into sentences and uses Semantic Chunking (Topic Shifts)."""
    def __init__(self, directory: str = ".", dense_model: Optional[CachedEmbeddingModel] = None):
        self.directory = directory
        self.model = dense_model or CachedEmbeddingModel("all-MiniLM-L6-v2")
        
    def parse_all(self) -> List[Dict[str, Any]]:
        all_chunks = []
        for root, dirs, files in os.walk(self.directory):
            for file in files:
                if file.endswith(".docx") and not file.startswith("~"):
                    file_path = os.path.join(root, file)
                    all_chunks.extend(self.parse_document(file_path))
        return all_chunks

    def _build_chunk(self, sentence_objs: List[Sentence], chunk_id: int, date: str, file_path: str, reason: str = "Document End") -> Dict[str, Any]:
        full_text = ""
        speakers = set()
        pages = set()
        
        last_speaker = None
        for s in sentence_objs:
            speakers.add(s.speaker)
            pages.add(s.page)
            if s.speaker != last_speaker:
                full_text += f"\n{s.speaker}: {s.text} "
                last_speaker = s.speaker
            else:
                full_text += f"{s.text} "
                
        page_list = sorted(list(pages))
        page_str = f"{page_list[0]}-{page_list[-1]}" if len(page_list) > 1 else (str(page_list[0]) if page_list else "1")
            
        return {
            "text": full_text.strip(),
            "speaker": ", ".join(list(speakers)[:3]) + ("..." if len(speakers) > 3 else ""),
            "date": date,
            "chunk_id": chunk_id,
            "page": page_str,
            "source_file": os.path.basename(file_path),
            "cut_reason": reason
        }

    def parse_document(self, path: str) -> List[Dict[str, Any]]:
        doc = docx.Document(path)
        date = self.extract_date(doc, os.path.basename(path))
        
        all_sentences = []
        current_speaker = "Unknown"
        current_page = 1
        page_char_count = 0
        
        for p in doc.paragraphs:
            lines = p.text.strip().split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                header_match = re.match(r"^([a-zA-Z\s\.]+)\s+\d{1,2}:\d{2}", line)
                if header_match and len(line.split()) < 15: 
                    current_speaker = header_match.group(1).strip()
                    continue
                
                # Split line into sentences
                sentences = re.split(r'(?<=[.!?])\s+', line)
                for s in sentences:
                    s = s.strip()
                    if s:
                        # Filter transcription noise
                        noise_phrases = ["stopped transcription", "started transcription", "joined the meeting", "left the meeting"]
                        if any(phrase in s.lower() for phrase in noise_phrases):
                            continue
                        all_sentences.append(Sentence(s, current_speaker, current_page, path))
                        page_char_count += len(s)
                
                if page_char_count > 2500:
                    current_page += 1
                    page_char_count = 0
                    
        chunks = []
        i = 0
        chunk_id = 0
        target_char_limit = 1200
        
        while i < len(all_sentences):
            current_chars = 0
            candidate_idx = i
            
            while candidate_idx < len(all_sentences) and current_chars < target_char_limit:
                current_chars += len(all_sentences[candidate_idx].text)
                candidate_idx += 1
                
            if candidate_idx == len(all_sentences):
                chunks.append(self._build_chunk(all_sentences[i:candidate_idx], chunk_id, date, path, "Document End"))
                break
                
            start_window = max(i + 1, candidate_idx - 30)
            end_window = min(len(all_sentences) - 1, candidate_idx + 30)
            
            if start_window >= end_window:
                best_cut = candidate_idx
                reason = "Target Size Fallback"
            else:
                window_sentences = all_sentences[start_window:end_window+1]
                texts = [s.text for s in window_sentences]
                embeddings = self.model.encode(texts, convert_to_tensor=True)
                min_sim = 1.0
                best_cut = candidate_idx
                
                for j in range(len(embeddings) - 1):
                    sim = util.cos_sim(embeddings[j], embeddings[j+1]).item()
                    if sim < min_sim:
                        min_sim = sim
                        best_cut = start_window + j + 1
                reason = f"Topic Shift Detected (Similarity: {min_sim:.3f})"
                
            chunks.append(self._build_chunk(all_sentences[i:best_cut], chunk_id, date, path, reason))
            i = best_cut
            chunk_id += 1
            
        return chunks

    def extract_date(self, doc, filename: str) -> str:
        date_pattern = r"\b(\d{1,4}[-/]\d{1,2}[-/]\d{1,4}|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})\b"
        for p in doc.paragraphs:
            match = re.search(date_pattern, p.text, re.IGNORECASE)
            if match:
                return match.group(1)
        match = re.search(date_pattern, filename, re.IGNORECASE)
        if match:
            return match.group(1)
        return "Unknown Date"


class VectorDatabase:
    """Shared Qdrant Vector Database Connector."""
    def __init__(self, collection_name: str = "teams_dense_collection"):
        self.qdrant_url = os.getenv("QDRANT_URL")
        self.qdrant_api_key = os.getenv("QDRANT_API_KEY")
        self.collection_name = collection_name
        
        if self.qdrant_url:
            self.client = QdrantClient(url=self.qdrant_url, api_key=self.qdrant_api_key)
        else:
            storage_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "qdrant_storage")
            os.makedirs(storage_path, exist_ok=True)
            self.client = QdrantClient(path=storage_path)
            
        self.dense_model = CachedEmbeddingModel("all-MiniLM-L6-v2")
        self.setup_collection()
        
    def setup_collection(self):
        try:
            if not self.client.collection_exists(self.collection_name):
                self.client.create_collection(
                    collection_name=self.collection_name,
                    vectors_config={
                        "dense": VectorParams(size=384, distance=Distance.COSINE)
                    }
                )
        except Exception:
            pass
            
    def insert_chunks(self, chunks: List[Dict[str, Any]]):
        points = []
        for c in chunks:
            dense_vec = self.dense_model.encode(c["text"])
            if hasattr(dense_vec, 'tolist'):
                dense_vec = dense_vec.tolist()
            
            points.append(PointStruct(
                id=str(uuid.uuid4()),
                payload=c,
                vector={"dense": dense_vec}
            ))
            
        if points:
            self.client.upsert(collection_name=self.collection_name, points=points)


_SHARED_DB = None
def get_vector_db() -> VectorDatabase:
    global _SHARED_DB
    if _SHARED_DB is None:
        _SHARED_DB = VectorDatabase()
    return _SHARED_DB


class CustomMeetingReranker:
    """Heuristic scoring reranker: Speaker match (+0.5), Date match (+0.5), Topic density (+0.05/keyword)."""
    def rerank(self, query: str, results: List[Any]) -> List[Any]:
        query_lower = query.lower()
        stop_words = {"what", "is", "the", "difference", "between", "an", "and", "according", "to", "did", "say", "about", "how", "are", "you"}
        keywords = [w for w in re.findall(r'\b\w+\b', query_lower) if w not in stop_words and len(w) > 3]
        
        for res in results:
            score = res.score if hasattr(res, 'score') else 1.0
            chunk_speaker = (res.payload or {}).get("speaker", "").lower()
            if any(s in query_lower and s in chunk_speaker for s in ["siddharth", "dakshinya", "himaya", "ganesh"]):
                score += 0.5
                
            chunk_date = (res.payload or {}).get("date", "").lower()
            if chunk_date and chunk_date != "unknown date" and chunk_date in query_lower:
                score += 0.5
                
            chunk_text = (res.payload or {}).get("text", "").lower()
            keyword_matches = sum(1 for k in keywords if k in chunk_text)
            score += (keyword_matches * 0.05)
            
            res.score = score
            
        return sorted(results, key=lambda x: x.score, reverse=True)


# ==============================================================================
# THE 4 RETRIEVAL PIPELINE IMPLEMENTATIONS
# ==============================================================================

def compute_batched_cosine_scores(query: str, chunks: List[Any], dense_model: CachedEmbeddingModel) -> List[Any]:
    """Computes cosine similarity between query and candidates in memory."""
    if not chunks:
        return []
    q_emb = dense_model.encode(query, convert_to_tensor=True)
    texts = [(c.payload or {}).get("text", "") for c in chunks]
    doc_embs = dense_model.encode(texts, convert_to_tensor=True)
    sims = util.cos_sim(q_emb, doc_embs)[0]
    
    for i, c in enumerate(chunks):
        c.score = float(sims[i].item())
    return chunks


def pipeline_p1_scroll_rerank(query: str, db: Optional[VectorDatabase] = None, top_per_doc: int = 2, total_candidates: int = 15) -> List[Any]:
    """
    P1 — Scroll + Custom Reranker
    Full database scroll scan -> per-document scoring -> top 2/doc -> top 15 candidates -> CustomMeetingReranker.
    """
    db = db or get_vector_db()
    raw_results, _ = db.client.scroll(collection_name=db.collection_name, limit=2000)
    scored = compute_batched_cosine_scores(query, raw_results, db.dense_model)
    
    # Document-balanced grouping
    by_doc = {}
    for p in scored:
        doc = (p.payload or {}).get("source_file", "unknown")
        by_doc.setdefault(doc, []).append(p)
        
    balanced_candidates = []
    for doc_file, plist in by_doc.items():
        plist.sort(key=lambda x: x.score, reverse=True)
        balanced_candidates.extend(plist[:top_per_doc])
        
    balanced_candidates.sort(key=lambda x: x.score, reverse=True)
    top_candidates = balanced_candidates[:total_candidates]
    
    reranker = CustomMeetingReranker()
    return reranker.rerank(query, top_candidates)


def pipeline_p2_scroll_scan(query: str, db: Optional[VectorDatabase] = None, top_per_doc: int = 4, total_candidates: int = 35) -> List[Any]:
    """
    P2 — Scroll API Scan
    Full database scroll scan -> per-document scoring -> top 4/doc -> top 35 candidates without reranker.
    """
    db = db or get_vector_db()
    raw_results, _ = db.client.scroll(collection_name=db.collection_name, limit=2000)
    scored = compute_batched_cosine_scores(query, raw_results, db.dense_model)
    
    by_doc = {}
    for p in scored:
        doc = (p.payload or {}).get("source_file", "unknown")
        by_doc.setdefault(doc, []).append(p)
        
    balanced_candidates = []
    for doc_file, plist in by_doc.items():
        plist.sort(key=lambda x: x.score, reverse=True)
        balanced_candidates.extend(plist[:top_per_doc])
        
    balanced_candidates.sort(key=lambda x: x.score, reverse=True)
    return balanced_candidates[:total_candidates]


def pipeline_p3_topk_vector(query: str, db: Optional[VectorDatabase] = None, top_k: int = 15) -> List[Any]:
    """
    P3 — Top-K Vector Search
    Native single-shot Qdrant query_points based purely on embedding cosine distance (fixed K=15).
    """
    db = db or get_vector_db()
    dense_vec = db.dense_model.encode(query)
    if hasattr(dense_vec, 'tolist'):
        dense_vec = dense_vec.tolist()
        
    return db.client.query_points(
        collection_name=db.collection_name,
        query=dense_vec,
        using="dense",
        limit=top_k
    ).points


def pipeline_p4_map_reduce(query: str, db: Optional[VectorDatabase] = None) -> Dict[str, List[Any]]:
    """
    P4 — Map-Reduce Retrieval Strategy
    Fetches 100% of chunks in Qdrant collection, grouped by source_file for two-stage Map summarization.
    """
    db = db or get_vector_db()
    raw_results, _ = db.client.scroll(collection_name=db.collection_name, limit=2000)
    grouped_by_doc = {}
    for p in raw_results:
        doc = (p.payload or {}).get("source_file", "unknown")
        grouped_by_doc.setdefault(doc, []).append(p)
    return grouped_by_doc


class DenseRetriever:
    """Standard Dense Retriever referencing CustomMeetingReranker."""
    def __init__(self, db: VectorDatabase):
        self.db = db
        self.reranker = CustomMeetingReranker()
        
    def retrieve(self, query: str, top_k: int = 10, rerank_top_k: int = 4):
        dense_vec = self.db.dense_model.encode(query)
        if hasattr(dense_vec, 'tolist'):
            dense_vec = dense_vec.tolist()
        results = self.db.client.query_points(
            collection_name=self.db.collection_name,
            query=dense_vec,
            using="dense",
            limit=top_k
        ).points
        if not results:
            return []
        results = self.reranker.rerank(query, results)
        return results[:rerank_top_k]
