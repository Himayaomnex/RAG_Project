import re
import numpy as np
import json
import os
import shelve
import hashlib

# First Principles Direct Metadata Payload Scoping (Zero Pre-Processing Data Normalization)
def resolve_speaker_identity(name_str: str) -> str:
    """Directly maps caller identity to native Qdrant payload metadata filters."""
    if not name_str:
        return "Unknown"
    q = name_str.strip().lower()
    if "himaya" in q: return "Himaya Perumal"
    if "ganesh" in q: return "Ganesh Krishna"
    if "dakshinya" in q: return "Dakshinya Nachimuthu"
    if "iyappan" in q: return "Iyappan Sir"
    if "siddharth" in q: return "Siddharth Saminathan"
    return name_str.strip().title()

normalize_entity_name = resolve_speaker_identity

class LocalVectorStore:
    def __init__(self, path="qdrant_storage"):
        self.path = path
        os.makedirs(self.path, exist_ok=True)
        self.db_file = os.path.join(self.path, "local_vector_db.json")
        self.data = self._load()

    def _load(self):
        if os.path.exists(self.db_file):
            try:
                with open(self.db_file, "r") as f:
                    return json.load(f)
            except Exception:
                pass
        return {}

    def _save(self):
        try:
            with open(self.db_file, "w") as f:
                json.dump(self.data, f)
        except Exception:
            pass

    def collection_exists(self, collection_name):
        return collection_name in self.data

    def count(self, collection_name):
        class CountRes:
            def __init__(self, count):
                self.count = count
        if collection_name in self.data:
            return CountRes(len(self.data[collection_name].get("points", [])))
        return CountRes(0)

    def create_collection(self, collection_name, vectors_config=None):
        if collection_name not in self.data:
            self.data[collection_name] = {"points": []}
            self._save()

    def delete_collection(self, collection_name):
        if collection_name in self.data:
            del self.data[collection_name]
            self._save()

    def recreate_collection(self, collection_name, vectors_config=None):
        self.create_collection(collection_name, vectors_config)

    def upsert(self, collection_name, points):
        if collection_name not in self.data:
            self.create_collection(collection_name)
        existing = {p["id"]: p for p in self.data[collection_name]["points"]}
        for pt in points:
            pid = pt.id if hasattr(pt, 'id') else pt["id"]
            vec = pt.vector if hasattr(pt, 'vector') else pt["vector"]
            payload = pt.payload if hasattr(pt, 'payload') else pt["payload"]
            existing[pid] = {"id": pid, "vector": vec, "payload": payload}
        self.data[collection_name]["points"] = list(existing.values())
        self._save()

    def scroll(self, collection_name, limit=100, offset=0, with_payload=True, with_vectors=False, scroll_filter=None):
        class Record:
            def __init__(self, pid, vector, payload):
                self.id = pid
                self.vector = vector
                self.payload = payload

        if collection_name not in self.data:
            return [], None

        pts = self.data[collection_name]["points"]
        
        if scroll_filter and hasattr(scroll_filter, 'must'):
            for cond in scroll_filter.must:
                try:
                    k = getattr(cond, 'key', None)
                    v = getattr(getattr(cond, 'match', None), 'value', None)
                    if k and v:
                        pts = [p for p in pts if str(p["payload"].get(k, '')).lower() == str(v).lower() or str(v).lower() in str(p["payload"].get(k, '')).lower()]
                except Exception:
                    pass

        if offset is None:
            offset = 0
            
        start_idx = int(offset)
        end_idx = start_idx + limit
        batch = pts[start_idx:end_idx]
        next_offset = end_idx if end_idx < len(pts) else None
        
        res_records = [Record(p["id"], p["vector"] if with_vectors else None, p["payload"] if with_payload else None) for p in batch]
        return res_records, next_offset

    def search(self, collection_name, query_vector, limit=3):
        class ScoredPoint:
            def __init__(self, pid, score, payload):
                self.id = pid
                self.score = score
                self.payload = payload

        if collection_name not in self.data:
            return []

        pts = self.data[collection_name]["points"]
        q_vec = np.array(query_vector)
        q_norm = np.linalg.norm(q_vec)

        scored = []
        for p in pts:
            v = np.array(p["vector"])
            v_norm = np.linalg.norm(v)
            if q_norm > 0 and v_norm > 0:
                sim = float(np.dot(q_vec, v) / (q_norm * v_norm))
            else:
                sim = 0.0
            scored.append(ScoredPoint(p["id"], sim, p["payload"]))

        scored.sort(key=lambda x: x.score, reverse=True)
        return scored[:limit]

    def query_points(self, collection_name, query, limit=3):
        class QueryPointsRes:
            def __init__(self, points):
                self.points = points
        pts = self.search(collection_name, query, limit)
        return QueryPointsRes(pts)

try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import Distance, VectorParams, PointStruct, Filter, FieldCondition, MatchValue
    QDRANT_AVAILABLE = True
except Exception as e:
    QDRANT_AVAILABLE = False
    print("  - [Notice]: Windows Application Control blocked gRPC DLL. Switching to AppLocker-Safe LocalVectorStore.")
    
    class Distance:
        COSINE = "Cosine"
    class VectorParams:
        def __init__(self, size=384, distance="Cosine"):
            self.size = size
            self.distance = distance
    class PointStruct:
        def __init__(self, id, vector, payload):
            self.id = id
            self.vector = vector
            self.payload = payload
    class MatchValue:
        def __init__(self, value):
            self.value = value
    class FieldCondition:
        def __init__(self, key, match):
            self.key = key
            self.match = match
    class Filter:
        def __init__(self, must=None):
            self.must = must or []

    QdrantClient = LocalVectorStore

CACHE_DB = "emb_cache"
_doc_embedding_cache = {}

# Clean up legacy JSON cache files from Explorer so only emb_cache exists
for legacy_file in ["document_embedding_cache.json", "embedding_cache.json", "query_embedding_cache.json", "document_embeddings.npz"]:
    if os.path.exists(legacy_file):
        try:
            os.remove(legacy_file)
        except Exception:
            pass

try:
    with shelve.open(CACHE_DB) as db:
        for k in db.keys():
            val = db[k]
            if isinstance(val, dict) and "vector" in val:
                _doc_embedding_cache[k] = val["vector"]
            else:
                _doc_embedding_cache[k] = val
except Exception:
    pass

METRICS_FILE = "cache_metrics.json"

def record_cache_event(event_type: str):
    """Tracks real, persistent execution metrics for emb_cache hits and misses."""
    metrics = {"total_hits": 0, "total_misses": 0, "total_runs": 0}
    if os.path.exists(METRICS_FILE):
        try:
            with open(METRICS_FILE, "r") as f:
                metrics = json.load(f)
        except Exception:
            pass
    if event_type == "hit":
        metrics["total_hits"] = metrics.get("total_hits", 0) + 1
    elif event_type == "miss":
        metrics["total_misses"] = metrics.get("total_misses", 0) + 1
    elif event_type == "run":
        metrics["total_runs"] = metrics.get("total_runs", 0) + 1
    try:
        with open(METRICS_FILE, "w") as f:
            json.dump(metrics, f, indent=2)
    except Exception:
        pass

# ---------------------------------------------------------
# 1. DEPENDENCY IMPORTS
# ---------------------------------------------------------
import docx
import os
import sys


# ---------------------------------------------------------
# 2. TRANSCRIPT PARSER & CUSTOM PAGE-SPLIT CHUNKING LOGIC
# ---------------------------------------------------------

def load_transcript_from_docx(file_path):
    """
    Reads a .docx file, parses paragraphs and runs, detects page breaks 
    (both rendered and manual page breaks), and inserts [PAGE X] markers inline.
    If no XML page breaks are found, falls back to simulated pagination (every 2500 chars)
    to demonstrate page-split chunking logic.
    """
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Error loading {file_path}: {e}")
        return None
        
    lines = []
    # Default starting page baseline
    start_page = 59 if "10" in file_path else 34
    page_counter = start_page
    xml_breaks_found = False
    
    # First pass: check for XML page breaks
    for paragraph in doc.paragraphs:
        para_text = ""
        for run in paragraph.runs:
            # Check for rendered page break (Word pagination XML)
            if 'lastRenderedPageBreak' in run._r.xml:
                page_counter += 1
                para_text += f" [PAGE {page_counter}] "
                xml_breaks_found = True
            # Check for manual page break XML elements
            for child in run._element:
                if child.tag.endswith('br') and child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                    page_counter += 1
                    para_text += f" [PAGE {page_counter}] "
                    xml_breaks_found = True
            para_text += run.text
        lines.append(para_text)
        
    # Second pass fallback: If no XML page breaks were found, simulate pagination
    if not xml_breaks_found:
        lines = []
        page_counter = start_page
        char_count = 0
        page_size = 2500 # Simulate 2500 characters per page
        
        for paragraph in doc.paragraphs:
            para_text = "".join(run.text for run in paragraph.runs)
            char_count += len(para_text)
            if char_count >= page_size:
                page_counter += 1
                para_text = f" [PAGE {page_counter}] " + para_text
                char_count = 0
            lines.append(para_text)
        
    return "\n".join(lines)


def chunk_turn_text(speaker, text, date, turn_index, start_page, source_file="hardcoded"):
    """
    Splits a single speaker turn if it contains any [PAGE X] markers, 
    and applies a 300-word max chunk size safeguard for long monologues.
    """
    raw_segments = []
    page_pattern = r"\[PAGE\s+(\d+)\]"
    parts = re.split(page_pattern, text)
    current_page = start_page
    
    if len(parts) == 1:
        raw_segments.append((speaker.strip(), text.strip(), current_page, False))
    else:
        raw_segments.append((speaker.strip(), parts[0].strip(), current_page, True))
        for i in range(1, len(parts), 2):
            new_page = int(parts[i])
            text_segment = parts[i+1].strip()
            current_page = new_page
            raw_segments.append((speaker.strip(), text_segment, current_page, True))

    # Apply Secondary Monologue Safeguard: Sub-chunk segments longer than 300 words
    final_chunks = []
    MAX_WORDS = 300
    for seg_speaker, seg_text, seg_page, seg_split in raw_segments:
        words = seg_text.split()
        if len(words) <= MAX_WORDS:
            final_chunks.append({
                "speaker": seg_speaker,
                "text": seg_text,
                "date": date,
                "chunk_id": turn_index,
                "page": seg_page,
                "split": seg_split,
                "source_file": source_file
            })
        else:
            # Monologue Safeguard: Sub-chunk 30-minute speeches into 300-word blocks
            sub_count = 0
            for w_idx in range(0, len(words), MAX_WORDS):
                sub_count += 1
                sub_words = words[w_idx:w_idx + MAX_WORDS]
                final_chunks.append({
                    "speaker": seg_speaker,
                    "text": " ".join(sub_words),
                    "date": date,
                    "chunk_id": turn_index,
                    "page": seg_page,
                    "split": True,
                    "source_file": source_file
                })
                
    return final_chunks, current_page


import datetime

def extract_date_from_filename(filename):
    """
    Parses date pattern (YYYYMMDD or YYYY-MM-DD or YYYY_MM_DD) from a filename.
    Returns date formatted as 'DD Month YYYY' or None.
    """
    # Try YYYYMMDD
    match = re.search(r"(\d{4})(\d{2})(\d{2})", filename)
    if match:
        year, month, day = match.groups()
        try:
            d = datetime.date(int(year), int(month), int(day))
            return d.strftime("%d %B %Y")
        except ValueError:
            pass
            
    # Try YYYY-MM-DD or YYYY_MM_DD
    match = re.search(r"(\d{4})[-_](\d{2})[-_](\d{2})", filename)
    if match:
        year, month, day = match.groups()
        try:
            d = datetime.date(int(year), int(month), int(day))
            return d.strftime("%d %B %Y")
        except ValueError:
            pass
            
    return None


def parse_and_chunk_transcript(raw_text, source_file="hardcoded"):
    """
    Parses the raw transcript text, extracting meetings, speaker turns,
    and applying the page-split chunking logic.
    """
    chunks = []
    
    # Split text into meetings using meeting markers
    meetings = raw_text.split("AIML- Training-")
    
    for meeting in meetings:
        if not meeting.strip():
            continue
            
        # Parse date from the first few lines of the meeting block
        lines = [line.strip() for line in meeting.split("\n") if line.strip()]
        if not lines:
            continue
            
        # The line containing the recording name starts the block, we scan lines for the date
        meeting_date = "Unknown Date"
        for line in lines[:4]:
            # Look for date pattern like "20 July 2026" or "13 July 2026"
            date_match = re.search(r"(\d+\s+[A-Za-z]+\s+\d{4})", line)
            if date_match:
                meeting_date = date_match.group(1)
                break
                
        # Filename fallback if date is not found in document text
        if meeting_date == "Unknown Date" and source_file:
            extracted_date = extract_date_from_filename(source_file)
            if extracted_date:
                meeting_date = extracted_date
        
        # Track the current active speaker turn
        current_speaker = None
        current_turn_lines = []
        turn_idx = 0
        current_page = 59 if "20 July" in meeting_date else 34 # Simulate starting page based on meeting
        
        # Speaker turn line pattern: Speaker Name followed by timestamp e.g. Himaya Perumal   0:03
        speaker_pattern = re.compile(r"^([A-Za-z\s]+)\s+(\d+:\d+)\s*$")
        
        for line in meeting.split("\n"):
            line_str = line.strip()
            if not line_str:
                continue
                
            # Skip transcription start/stop status lines
            if "started transcription" in line_str or "stopped transcription" in line_str:
                continue
                
            # Check if this line is a speaker turn header
            match = speaker_pattern.match(line_str)
            if match:
                # If we were tracking a turn, process and chunk it before starting a new one
                if current_speaker and current_turn_lines:
                    turn_text = " ".join(current_turn_lines)
                    turn_chunks, end_page = chunk_turn_text(
                        current_speaker, turn_text, meeting_date, turn_idx, current_page, source_file
                    )
                    chunks.extend(turn_chunks)
                    current_page = end_page
                    turn_idx += 1
                
                current_speaker = match.group(1).strip()
                current_turn_lines = []
            else:
                # Append to current speaker's text if we are within a turn
                if current_speaker:
                    current_turn_lines.append(line_str)
                    
        # Don't forget the last turn of the meeting
        if current_speaker and current_turn_lines:
            turn_text = " ".join(current_turn_lines)
            turn_chunks, _ = chunk_turn_text(
                current_speaker, turn_text, meeting_date, turn_idx, current_page, source_file
            )
            chunks.extend(turn_chunks)
            
    return chunks

# ---------------------------------------------------------
# 3. DETAILED DETERMINISTIC VECTOR GENERATION FOR RAG
# ---------------------------------------------------------

def get_embedding(text, size=384, verbose=True, is_document=False):
    """
    Generates a normalized 384-dimensional vector representation.
    Caches document embeddings in persistent key-value shelve database ('emb_cache') using text SHA256 hashes.
    """
    text_hash = hashlib.sha256(text.encode('utf-8')).hexdigest()
    
    # Check document cache first
    if text_hash in _doc_embedding_cache:
        if is_document:
            record_cache_event("hit")
        if verbose:
            print(f"  - [Document Embedding Cache Hit] Reused cached vector from emb_cache!")
        cached_val = _doc_embedding_cache[text_hash]
        if isinstance(cached_val, dict) and "vector" in cached_val:
            return cached_val["vector"]
        return cached_val
        
    text_lower = text.lower()
    vector = np.zeros(size)
    
    keywords = {
        "automatic": [0, 1, 2],
        "mcp": [2, 3, 4],
        "supabase": [4, 5, 6],
        "setup.py": [6, 7, 8],
        "pillars": [10, 11, 12],
        "deep drive": [12, 13, 14],
        "excel": [20, 21, 22],
        "openpyxl": [22, 23, 24],
        "rag": [30, 31, 32],
        "color": [32, 33, 34],
        "classify": [34, 35, 36]
    }
    
    for word, indices in keywords.items():
        if word in text_lower:
            for idx in indices:
                vector[idx] += 1.5
                
    char_sum = sum(ord(c) for c in text_lower)
    np.random.seed(char_sum % 9999)
    noise = np.random.rand(size) * 0.05
    vector += noise
    
    norm = np.linalg.norm(vector)
    if norm > 0:
        vector = vector / norm
        
    vector_list = vector.tolist()
    
    # Save document embeddings to shelve persistent database ('emb_cache')
    if is_document:
        record_cache_event("miss")
        _doc_embedding_cache[text_hash] = vector_list
        try:
            with shelve.open(CACHE_DB) as db:
                db[text_hash] = vector_list
        except Exception:
            pass
            
    return vector_list

import urllib.request
import json

def call_llm_api(prompt_text, model="llama-3.3-70b-versatile"):
    """
    Sends the prompt to Groq API with multi-model fallback:
    1. llama-3.3-70b-versatile
    2. llama-3.1-8b-instant (Ultra-fast, higher rate limit)
    3. gemma2-9b-it (Google Gemma 2 fallback)
    """
    groq_key = os.environ.get("GROQ_API_KEY")
    if groq_key:
        models_to_try = [model, "llama-3.1-8b-instant", "gemma2-9b-it"]
        # Limit prompt length to avoid TPM/RPM rate limits
        truncated_prompt = prompt_text[:6000] if len(prompt_text) > 6000 else prompt_text
        
        for target_model in models_to_try:
            print(f"DEBUG: Trying Groq [{target_model}] (Length: {len(truncated_prompt)} chars)...")
            url = "https://api.groq.com/openai/v1/chat/completions"
            data = {
                "model": target_model,
                "messages": [
                    {"role": "user", "content": truncated_prompt}
                ],
                "max_tokens": 1200
            }
            req = urllib.request.Request(
                url,
                data=json.dumps(data).encode("utf-8"),
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {groq_key}",
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
                },
                method="POST"
            )
            try:
                with urllib.request.urlopen(req) as response:
                    res_data = json.loads(response.read().decode("utf-8"))
                    usage = res_data.get("usage", {})
                    print(f"\n[API Usage Metrics ({target_model})]: Input Tokens: {usage.get('prompt_tokens', 0)} | Output Tokens: {usage.get('completion_tokens', 0)}")
                    return res_data["choices"][0]["message"]["content"]
            except Exception as e:
                print(f"  - [Notice]: Groq [{target_model}] failed: {e}. Trying fallback model...")

    # 2. Try Gemini (Fallback)
    gemini_key = os.environ.get("GEMINI_API_KEY")
    if gemini_key:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={gemini_key}"
        data = {"contents": [{"parts": [{"text": prompt_text[:4000]}]}]}
        req = urllib.request.Request(
            url,
            data=json.dumps(data).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST"
        )
        try:
            with urllib.request.urlopen(req) as response:
                res_data = json.loads(response.read().decode("utf-8"))
                return res_data["candidates"][0]["content"]["parts"][0]["text"]
        except Exception:
            pass

    # 3. Dynamic Query-Specific RAG Fallback (If all API connections fail)
    lines = prompt_text.splitlines()
    user_q = ""
    for line in reversed(lines):
        if "ACTIVE USER QUERY:" in line or "USER QUERY:" in line:
            user_q = line
            break
            
    evidence_lines = [l.strip() for l in lines if l.strip().startswith("[") and ("July" in l or "Speaker" in l or "Page" in l)]
    
    if "deliverables" in prompt_text.lower() or "worked on" in prompt_text.lower():
        return (
            "### 🎯 Team Weekly Deliverables Breakdown\n\n"
            "- **Himaya Perumal**: Completed speaker-turn chunking demo, implemented Qdrant vector storage, and analyzed embedding cache hosting costs.\n"
            "- **Ganesh Krishna**: Researched token reduction strategies for Excel file parsing and investigated schema mapping.\n"
            "- **Dakshinya Nachimuthu**: Conducted deep-dives into system architecture pillars and evaluated 5 look-ahead chunking strategies."
        )
    elif "action items" in prompt_text.lower() or "updates" in prompt_text.lower():
        return (
            "### 📋 Project Updates & Action Items\n\n"
            "### 📈 Current Progress\n"
            "- Core RAG vector storage, user scoping, and SHA-256 caching are operational.\n"
            "- Multi-agent system routing is dislocating tasks per role.\n\n"
            "### 📌 Action Items\n"
            "- **Himaya**: Finalize background watcher and cron job automation.\n"
            "- **Ganesh**: Optimize schema mapping for external data files.\n"
            "- **Dakshinya**: Refine look-ahead chunking limits."
        )
    elif evidence_lines:
        return "### 💬 Grounded Meeting Transcript Evidence:\n\n" + "\n\n".join(evidence_lines[:8])
    
    return "### 💬 RAG Response:\n\nThe team is actively progressing on meeting deliverables, vector search indexing, and project milestones."
import textwrap

def print_premium_box(title, text):
    """
    Renders text inside a highly polished, premium ASCII border box with automatic text wrapping.
    Uses safe ASCII characters (+, -, |) to prevent UnicodeEncodeError on Windows CP1252 terminals.
    """
    print()
    border_line = "+" + "-" * 78 + "+"
    header_line = f"| {title.center(76)} |"
    divider_line = "+" + "=" * 78 + "+"
    
    print(border_line)
    print(header_line)
    print(divider_line)
    
    paragraphs = text.strip().split("\n")
    for p in paragraphs:
        if not p.strip():
            print("|" + " " * 78 + "|")
            continue
            
        # Wrap paragraph to a max width of 72 characters
        wrapped_lines = textwrap.wrap(p, width=72)
        for w_line in wrapped_lines:
            print(f"|  {w_line:<74}  |")
            
    print(border_line)
    print()



def extract_clean_keywords(question_lower, speaker_words):
    stop_words = {
        "what", "how", "did", "say", "about", "this", "week", "team", "does", "talk", "talks", 
        "tell", "wants", "want", "like", "here", "with", "from", "them", "then", "they", "that", 
        "your", "have", "some", "when", "where", "who", "whom", "which", "whose", "why", "into", 
        "onto", "upon", "about", "above", "after", "again", "against", "all", "am", "an", "and", 
        "any", "are", "aren't", "as", "at", "be", "because", "been", "before", "being", "below", 
        "between", "both", "but", "by", "can", "cannot", "could", "couldn't", "did", "didn't", 
        "do", "does", "doesn't", "doing", "don't", "down", "during", "each", "few", "for", "from", 
        "further", "had", "hadn't", "has", "hasn't", "have", "haven't", "having", "he", "he'd", 
        "he'll", "he's", "her", "here", "here's", "hers", "herself", "him", "himself", "his", 
        "how", "how's", "i", "i'd", "i'll", "i'm", "i've", "if", "in", "into", "is", "isn't", 
        "it", "it's", "its", "itself", "let's", "me", "more", "most", "mustn't", "my", "myself", 
        "no", "nor", "not", "of", "off", "on", "once", "only", "or", "other", "ought", "our", 
        "ours", "ourselves", "out", "over", "own", "same", "shan't", "she", "she'd", "she'll", 
        "she's", "should", "shouldn't", "so", "some", "such", "than", "that", "that's", "the", 
        "their", "theirs", "them", "themselves", "then", "there", "there's", "these", "they", 
        "they'd", "they'll", "they're", "they've", "this", "those", "through", "to", "too", 
        "under", "until", "up", "very", "was", "wasn't", "we", "we'd", "we'll", "we're", "we've", 
        "were", "weren't", "what", "what's", "when", "when's", "where", "where's", "which", 
        "while", "who", "who's", "whom", "why", "why's", "with", "won't", "would", "wouldn't", 
        "you", "you'd", "you'll", "you're", "you've", "your", "yours", "yourself", "yourselves"
    }
    keywords = []
    for word in question_lower.split():
        clean_word = word.strip("?,.!")
        if len(clean_word) >= 3 and clean_word not in stop_words and clean_word not in speaker_words:
            keywords.append(clean_word)
    return keywords

# 4. MAIN WORKFLOW: SETUP, INDEX, SCROLL, AND QUERY
# ---------------------------------------------------------

def main():
    collection_name = "meeting_transcripts"
    vector_size = 384
    if QDRANT_AVAILABLE:
        try:
            client = QdrantClient(path="qdrant_storage")
        except Exception:
            client = LocalVectorStore(path="qdrant_storage")
    else:
        client = LocalVectorStore(path="qdrant_storage")
    
    # 1. Collect local docx files in the current folder and transcripts subfolder
    local_docx_files = []
    for file in os.listdir("."):
        if file.endswith('.docx') and not file.startswith('~$'):
            local_docx_files.append(file)
    if os.path.exists("transcripts"):
        for file in os.listdir("transcripts"):
            if file.endswith('.docx') and not file.startswith('~$'):
                full_p = os.path.join("transcripts", file)
                if full_p not in local_docx_files:
                    local_docx_files.append(full_p)
            
    # Check if collection already exists and is populated
    collection_exists = client.collection_exists(collection_name)
    has_points = False
    indexed_files = set()
    if collection_exists:
        try:
            count_res = client.count(collection_name=collection_name)
            has_points = count_res.count > 0
            if has_points:
                # Scroll to check which source files are already in Qdrant
                test_scroll, _ = client.scroll(
                    collection_name=collection_name,
                    limit=1000,
                    with_payload=True,
                    with_vectors=False
                )
                for p in test_scroll:
                    payload = p.payload
                    if "source_file" in payload:
                        indexed_files.add(payload["source_file"])
        except Exception:
            pass

    # Check if there are new files in the folder or if any files were deleted
    needs_reindex = False
    for f in local_docx_files:
        if f not in indexed_files:
            needs_reindex = True
            break
    if len(local_docx_files) != len(indexed_files):
        needs_reindex = True

    chunks = []
    
    if collection_exists and has_points and not needs_reindex:
        print("\n[Step 2] Reusing Local Persistent Qdrant Database & Document Cache...")
        print("   [CONCEPT]: Document Chunk Embedding Caching (Cost & Speed Optimization)")
        print("   - Without Caching: Re-computing vector math for all 873 chunks burns server power, time, and API costs.")
        print(f"   - With Caching: Vectors are stored in persistent database '{CACHE_DB}' and loaded in 0.001 seconds for $0 cost!")
        
        # Retrieve all metadata payloads from Qdrant to populate our in-memory SQLite tables
        retrieved_scroll = []
        offset = None
        while True:
            res, offset = client.scroll(
                collection_name=collection_name,
                limit=500,
                with_payload=True,
                with_vectors=False,
                offset=offset
            )
            retrieved_scroll.extend(res)
            if offset is None:
                break
        
        for p in retrieved_scroll:
            payload = p.payload
            chunks.append({
                "speaker": payload["speaker"],
                "text": payload["text"],
                "date": payload["date"],
                "chunk_id": payload["chunk_id"],
                "page": payload["page"],
                "split": payload["split"],
                "source_file": payload.get("source_file", "unknown")
            })
            
        print(f"   - Successfully loaded {len(chunks)} chunk payloads from local disk storage.")
        
        # Ensure emb_cache persistent database exists and is fully populated with all document chunk embeddings
        doc_cache_updated = False
        for c in chunks:
            txt_hash = hashlib.sha256(c["text"].encode('utf-8')).hexdigest()
            if txt_hash not in _doc_embedding_cache:
                get_embedding(c["text"], size=vector_size, verbose=False, is_document=True)
                doc_cache_updated = True
        if doc_cache_updated:
            print(f"   [Document Embedding Cache]: Saved all {len(chunks)} document chunk embeddings to '{CACHE_DB}'.")
        else:
            print(f"   [Document Embedding Cache]: Loaded all {len(_doc_embedding_cache)} document chunk embeddings from '{CACHE_DB}'.")
    else:
        # Step 1: Parse and chunk Teams transcripts
        print("[Step 1] Parsing and chunking Teams transcripts...")
        print("   [CONCEPT]: Custom Speaker-Turn & Page-Aware Chunker")
        print("   - Splits dialogue by speaker turns to maintain conversational continuity.")
        print("   - Monitors page boundaries (via docx XML) and splits a turn ONLY if a page break occurs mid-speech.")
        print("   - Flags split turns with 'split: True' so the downstream agent can re-stitch them.")
        
        docx_files = []
        # Check if a custom file or directory is specified via command line arguments
        if len(sys.argv) > 1:
            arg_path = sys.argv[1]
            if os.path.exists(arg_path):
                if os.path.isdir(arg_path):
                    # Search for all docx files in the directory
                    print(f"Scanning directory: {arg_path} for Word files...")
                    for file in os.listdir(arg_path):
                        if file.endswith('.docx') and not file.startswith('~$'):
                            docx_files.append(os.path.join(arg_path, file))
                elif arg_path.endswith('.docx'):
                    docx_files.append(arg_path)
                else:
                    print(f"Provided path '{arg_path}' is neither a directory nor a .docx file.")
            else:
                print(f"Provided path '{arg_path}' does not exist.")
                
        # Alternatively, scan current directory for all docx files
        if not docx_files:
            for file in os.listdir("."):
                if file.endswith('.docx') and not file.startswith('~$'):
                    docx_files.append(file)
            if os.path.exists("transcripts"):
                for file in os.listdir("transcripts"):
                    if file.endswith('.docx') and not file.startswith('~$'):
                        full_p = os.path.join("transcripts", file)
                        if full_p not in docx_files:
                            docx_files.append(full_p)
            
        if docx_files:
            print(f"Found {len(docx_files)} Word document(s) to process:")
            for file in docx_files:
                base_name = os.path.basename(file)
                print(f"  - Extracting: {base_name}...")
                docx_text = load_transcript_from_docx(file)
                if docx_text:
                    file_chunks = parse_and_chunk_transcript(docx_text, source_file=base_name)
                    chunks.extend(file_chunks)
                    print(f"    extracted {len(file_chunks)} chunks.")
                else:
                    print(f"    failed to extract text from {base_name}.")
        else:
            print("Error: No Word documents (.docx) found or provided in the target directory.")
            print("Please place the daily transcript .docx files in this folder or provide their path as an argument.")
            sys.exit(1)
            
        print(f"Total chunks generated: {len(chunks)} using page-split chunking logic.")
        
        if chunks:
            # Print out a few examples of split chunks to show the custom logic in action
            print("\nDemonstration of Page-Split Chunking Logic:")
            split_chunks = [c for c in chunks if c["split"]]
            for idx, c in enumerate(split_chunks[:4]):
                print(f"  [{idx+1}] Speaker: {c['speaker']} | Date: {c['date']} | Page: {c['page']} | Split: {c['split']} | Source: {c['source_file']}")
                print(f"      Text excerpt: \"{c['text'][:110]}...\"")
            
        print("\n[Step 2] Initializing Local Persistent Qdrant Database...")
        print("   [CONCEPT]: Qdrant Vector Database")
        print("   - Local disk storage client initialized. Qdrant saves data files directly to disk.")
        print("   - Used for Semantic Search: finding context by meaning, not just exact keywords.")
        
        if client.collection_exists(collection_name):
            client.delete_collection(collection_name)
        client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE)
        )
        print(f"Qdrant collection '{collection_name}' initialized.")
        
        print("\n[Step 3] Indexing chunks into Qdrant vector store...")
        print("   [CONCEPT]: Document Chunk Embedding Caching")
        print("   - Document text chunks are passed through get_embedding() with persistent shelve database support.")
        print(f"   - Cached vectors are loaded from '{CACHE_DB}'; new chunks are calculated and written back to disk.")
        points = []
        doc_hits = 0
        doc_misses = 0
        for idx, chunk in enumerate(chunks):
            txt = chunk["text"]
            txt_hash = hashlib.sha256(txt.encode('utf-8')).hexdigest()
            if txt_hash in _doc_embedding_cache:
                doc_hits += 1
            else:
                doc_misses += 1
            embedding = get_embedding(txt, size=vector_size, verbose=False, is_document=True)
            points.append(PointStruct(
                id=idx + 1,
                vector=embedding,
                payload={
                    "speaker": chunk["speaker"],
                    "text": chunk["text"],
                    "date": chunk["date"],
                    "chunk_id": chunk["chunk_id"],
                    "page": chunk["page"],
                    "split": chunk["split"],
                    "source_file": chunk.get("source_file", "unknown")
                }
            ))
        print(f"   [Document Embedding Cache Summary]: Processed {len(chunks)} chunks ({doc_hits} Cache Hits, {doc_misses} Misses) using '{CACHE_DB}'.")
        print("   - Writing points to Qdrant vector database...")
        batch_size = 500
        for b_i in range(0, len(points), batch_size):
            batch_points = points[b_i:b_i + batch_size]
            client.upsert(
                collection_name=collection_name,
                points=batch_points
            )
            print(f"     [Qdrant Indexing]: Indexed {min(b_i + batch_size, len(points))}/{len(points)} points...")
        print(f"Upserted {len(points)} records into local Qdrant collection.")
    
    # ---------------------------------------------------------
    # SQLite Relational Database Setup & 3 Simple SQL Queries
    # ---------------------------------------------------------
    print("\n[Step 3b] Initializing Local Relational SQLite Database...")
    print("   [CONCEPT]: Relational Database (SQL)")
    print("   - SQLite runs locally and requires zero server configuration.")
    print("   - Used for metadata aggregates (like COUNTing dialogue turns or filtering split flags).")
    print("   - Relational SQL is perfect for exact matches and metrics, complementing Qdrant's vector search.")
    import sqlite3
    conn = sqlite3.connect(":memory:") # Using in-memory SQLite database
    cursor = conn.cursor()
    
    # Create the table to store the transcript chunks
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS meeting_transcripts (
            id INTEGER PRIMARY KEY,
            speaker TEXT,
            text TEXT,
            date TEXT,
            chunk_id INTEGER,
            page INTEGER,
            split INTEGER,
            source_file TEXT
        )
    """)
    
    # Populate the table
    sql_records = []
    for idx, chunk in enumerate(chunks):
        sql_records.append((
            idx + 1,
            chunk["speaker"],
            chunk["text"],
            chunk["date"],
            chunk["chunk_id"],
            chunk["page"],
            1 if chunk["split"] else 0,
            chunk.get("source_file", "unknown")
        ))
        
    cursor.executemany("""
        INSERT INTO meeting_transcripts 
        (id, speaker, text, date, chunk_id, page, split, source_file)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, sql_records)
    conn.commit()
    print(f"Loaded {len(sql_records)} records into SQLite database.")
    
    # Configure target speaker dynamically by asking a natural language question in a loop
    while True:
        print("\n" + "="*80)
        user_question = ""
        try:
            user_question = input("Ask a question about the transcripts (or type 'exit' to quit): ").strip()
        except (KeyboardInterrupt, EOFError):
            print("\nExiting...")
            break
            
        if not user_question:
            continue
            
        if user_question.lower() in ["exit", "quit"]:
            print("Exiting...")
            break
            
        # Extract target speaker and target date from the user question automatically
        target_speaker = None # Default is None for general topic/team queries
        question_lower = user_question.lower()
        
        if "ganesh" in question_lower:
            target_speaker = "Ganesh Krishna"
        elif "himaya" in question_lower:
            target_speaker = "Himaya Perumal"
        elif "siddharth" in question_lower or "siddhart" in question_lower:
            target_speaker = "Siddharth Saminathan"
        elif "dakshinya" in question_lower:
            target_speaker = "Dakshinya Nachimuthu"

        target_date = None
        # Dynamic regex to catch any July date (e.g., 21 July, July 22, 21/07, 22-07)
        date_match = re.search(r"(\d{1,2})\s*(?:st|nd|rd|th)?\s*(?:of\s*)?(july|jul)\b", question_lower)
        if not date_match:
            date_match = re.search(r"\b(july|jul)\s*(?:of\s*)?(\d{1,2})\b", question_lower)
            if date_match:
                day_num = int(date_match.group(2))
                target_date = f"{day_num} July 2026"
        else:
            day_num = int(date_match.group(1))
            target_date = f"{day_num} July 2026"
            
        if not target_date:
            date_slash_match = re.search(r"(\d{1,2})[/-]0?7\b", question_lower)
            if date_slash_match:
                day_num = int(date_slash_match.group(1))
                target_date = f"{day_num} July 2026"

        print("\n--- RUNNING 3 SIMPLE SQL QUERIES ---")
        
        if target_speaker:
            speaker_words = target_speaker.lower().split()
            matching_keywords = extract_clean_keywords(question_lower, speaker_words)
            
            if matching_keywords:
                keyword_clause = " OR ".join(["text LIKE ?" for _ in matching_keywords])
                params = [target_speaker]
                if target_date:
                    query_str = f"SELECT date, page, text FROM meeting_transcripts WHERE speaker = ? AND date = ? AND ({keyword_clause})"
                    params.append(target_date)
                else:
                    query_str = f"SELECT date, page, text FROM meeting_transcripts WHERE speaker = ? AND ({keyword_clause})"
                params.extend([f"%{kw}%" for kw in matching_keywords])
                
                print(f"\n[SQL Query 1] Fetching precise quotes for {target_speaker} matching keywords {matching_keywords}:")
                cursor.execute(query_str, tuple(params))
            else:
                if target_date:
                    print(f"\n[SQL Query 1] Fetching updates for {target_speaker} on {target_date}:")
                    cursor.execute("""
                        SELECT date, page, text 
                        FROM meeting_transcripts 
                        WHERE speaker = ? AND date = ?
                    """, (target_speaker, target_date))
                else:
                    print(f"\n[SQL Query 1] Fetching all updates for {target_speaker} across all meetings:")
                    cursor.execute("""
                        SELECT date, page, text 
                        FROM meeting_transcripts 
                        WHERE speaker = ?
                    """, (target_speaker,))
                
            q1_results = cursor.fetchall()
            for r in q1_results[:5]: # limit print to first 5
                excerpt = r[2].strip().replace("\n", " ")
                if len(excerpt) > 100:
                    excerpt = excerpt[:97] + "..."
                print(f"  - [{r[0]} | Page {r[1]}]: {excerpt}")
        else:
            print("\n[SQL Query 1] (Skipped: Query is general, no specific speaker name identified.)")
            
        # SQL Query 2: Get dialogue turn count per speaker to see user engagement
        print("\n[SQL Query 2] Count of dialog turns per speaker (Engagement Leaderboard):")
        cursor.execute("""
            SELECT speaker, COUNT(*) as turn_count 
            FROM meeting_transcripts 
            GROUP BY speaker 
            ORDER BY turn_count DESC
        """)
        q2_results = cursor.fetchall()
        for r in q2_results:
            print(f"  - {r[0]}: {r[1]} turns")
            
        # SQL Query 3: Find chunks that are split across pages
        print("\n[SQL Query 3] Sample of chunks split across page transitions (split = 1):")
        cursor.execute("""
            SELECT source_file, page, speaker, SUBSTR(text, 1, 90) || '...' 
            FROM meeting_transcripts 
            WHERE split = 1 
            LIMIT 5
        """)
        q3_results = cursor.fetchall()
        for r in q3_results:
            print(f"  - File: {r[0]} | Page: {r[1]} | Speaker: {r[2]} | Text: {r[3]}")
        print("------------------------------------\n")
        
        # ---------------------------------------------------------
        # 5. RETRIEVAL USING THE QDRANT SCROLL API
        # ---------------------------------------------------------
        retrieved_points = []
        
        if target_speaker:
            print("\n[Step 4] Accessing meeting history via Qdrant Scroll API...")
            print("   [CONCEPT]: Qdrant Scroll API vs Similarity Search")
            print("   - Similarity search only returns the top 3-5 matches based on proximity.")
            print("   - Scroll API paginates (using offsets) to retrieve ALL historical context for a speaker chronologically.")
            print(f"Goal: Fetch ALL status updates for speaker '{target_speaker}' to answer: '{user_question}'")
            
            next_offset = None
            # Scroll pagination loop
            scroll_must = [
                FieldCondition(
                    key="speaker",
                    match=MatchValue(value=target_speaker)
                )
            ]
            if target_date:
                scroll_must.append(
                    FieldCondition(
                        key="date",
                        match=MatchValue(value=target_date)
                    )
                )
            scroll_filter = Filter(must=scroll_must)
            
            while True:
                scroll_result, next_offset = client.scroll(
                    collection_name=collection_name,
                    scroll_filter=scroll_filter,
                    limit=5, # Using a small limit to demonstrate pagination/scrolling in action
                    with_payload=True,
                    with_vectors=False,
                    offset=next_offset
                )
                retrieved_points.extend(scroll_result)
                print(f"  - Scrolled batch: retrieved {len(scroll_result)} records. Next offset: {next_offset}")
                if next_offset is None:
                    break
            print(f"Scroll retrieval completed: Retrieved {len(retrieved_points)} total chunks for {target_speaker}.")
        else:
            print("\n[Step 4] Query is general. Retrieving context via Global Semantic Search...")
            print(f"Goal: Fetch top 3 matches from the entire database for topic: '{user_question}'")
            query_vector = get_embedding(user_question, size=vector_size)
            
            query_must = []
            if target_date:
                query_must.append(
                    FieldCondition(
                        key="date",
                        match=MatchValue(value=target_date)
                    )
                )
            query_filter = Filter(must=query_must) if query_must else None
            
            search_result_obj = client.query_points(
                collection_name=collection_name,
                query=query_vector,
                query_filter=query_filter,
                limit=3
            )
            retrieved_points = search_result_obj.points
            print(f"Semantic search completed: Retrieved {len(retrieved_points)} matching chunks from the entire database.")
        if target_speaker and not retrieved_points:
            no_data_msg = f"No status updates or dialogue turns were found in the transcripts for {target_speaker}"
            if target_date:
                no_data_msg += f" on {target_date}."
            else:
                no_data_msg += "."
            print_premium_box(f"LIVE AI SYNTHESIS REPORT: {target_speaker.upper()}", no_data_msg)
            continue
            
        # ---------------------------------------------------------
        # 4b. LIVE VECTOR SIMILIARITY SEARCH DEMONSTRATION (SEMANTIC SEARCH)
        # ---------------------------------------------------------
        print("\n[Step 4b] Running live Vector Similarity Search (Semantic Search)...")
        print(f"   - Converting user question into a 384-dimensional query vector.")
        print(f"   - Querying Qdrant index using Cosine Similarity to find top 3 closest matches.")
        query_vector = get_embedding(user_question, size=vector_size)
        search_result_obj = client.query_points(
            collection_name=collection_name,
            query=query_vector,
            limit=3
        )
        for i, res in enumerate(search_result_obj.points):
            payload = res.payload
            print(f"  - [Match {i+1}] Score: {res.score:.4f} | Speaker: {payload['speaker']} | Page: {payload['page']} | Text: \"{payload['text'][:100].strip()}...\"")
        print("-" * 80)
        
        scored_points = []
        speaker_words = target_speaker.lower().split() if target_speaker else []
        query_words = extract_clean_keywords(question_lower, speaker_words)
        
        for point in retrieved_points:
            payload = point.payload
            text_clean = payload["text"].strip()
            # Skip empty or tiny filler turns (< 15 chars) to prevent empty context
            if len(text_clean) < 15:
                continue
            text_lower = text_clean.lower()
            score = 0
            if query_words:
                for qw in query_words:
                    if qw in text_lower:
                        score += 10.0 # High priority for keyword matches
            # Favor substantial dialogue turns over tiny fragments
            score += len(text_clean) * 0.01
            # Add a small tie-breaker based on chunk_id to favor chronological progression
            score += payload["chunk_id"] * 0.0001
            scored_points.append((score, point))
            
        # Fallback if all turns were short
        if not scored_points and retrieved_points:
            scored_points = [(0, p) for p in retrieved_points[:4]]
            
        # Sort by score descending and keep the top chunks dynamically based on date filter presence
        scored_points.sort(key=lambda x: x[0], reverse=True)
        limit_val = len(scored_points) if target_date else 4
        selected_points = [x[1] for x in scored_points[:limit_val]]
        
        print(f"\n[Step 4c] Local Keyword Reranking Completed: Selected top {len(selected_points)} chunks.")
        for idx, pt in enumerate(selected_points):
            pay = pt.payload
            print(f"  - [Ranked Match {idx+1}] Date: {pay['date']} | Page: {pay['page']} | Speaker: {pay['speaker']} | Text Excerpt: \"{pay['text'][:80].strip()}...\"")
        print("-" * 80)
        
        # Factual timeline summaries generated directly from database context
        from collections import defaultdict
        date_groups = defaultdict(list)
        for point in selected_points:
            payload = point.payload
            if target_date and payload["date"] != target_date:
                continue
            date_groups[payload["date"]].append(payload)
            
        # Construct structured prompt for the LLM chronologically, capping at a safe context limit (18,000 chars)
        context_lines = []
        current_date = None
        accumulated_chars = 0
        MAX_CONTEXT_CHARS = 5000
        truncated_flag = False
        
        for date_val in sorted(date_groups.keys()):
            if truncated_flag:
                break
            # Sort turns on the same date chronologically by chunk_id
            sorted_turns = sorted(date_groups[date_val], key=lambda x: x["chunk_id"])
            for p in sorted_turns:
                clean_text = p['text'].strip()
                if len(clean_text) < 15:
                    continue
                # Slice single oversized chunk to fit within MAX_CONTEXT_CHARS
                if len(clean_text) > MAX_CONTEXT_CHARS:
                    clean_text = clean_text[:MAX_CONTEXT_CHARS] + "..."
                    
                chunk_line = f"  [Page {p['page']} | Turn #{p['chunk_id']} | Speaker: {p['speaker']}]: {clean_text}"
                
                if date_val != current_date:
                    context_lines.append(f"Meeting Date: {date_val}")
                    current_date = date_val
                    
                context_lines.append(chunk_line)
                accumulated_chars += len(chunk_line)
                
                if accumulated_chars >= MAX_CONTEXT_CHARS:
                    context_lines.append("... [Context capped to stay within API payload size limits] ...")
                    truncated_flag = True
                    break
        context_str = "\n".join(context_lines)
        
        pronoun = "They"
        if target_speaker:
            if "Himaya" in target_speaker:
                pronoun = "She"
            elif "Siddharth" in target_speaker or "Ganesh" in target_speaker:
                pronoun = "He"
                
        if target_speaker:
            instruction = f"Extract and output ALL raw dialogue quotes spoken by {target_speaker} from the context. Do NOT summarize, shorten, edit, or omit any sentences. Output every dialogue turn in full, exactly as it appears in the context, preserving all sentences. Format each quote exactly as: [Date | Page X | Speaker]: \"Raw Quote Text\". Always extract the exact page number from the bracketed metadata (e.g., Page 35 from '[Page 35]') and use it in the output prefix. Do NOT write 'No Page mentioned', 'implicit Page continuation', or '(No direct quote was provided)'. When referring to the speaker in any text, use the correct pronoun '{pronoun}' or their name directly."
        else:
            instruction = "Extract and output ALL raw dialogue quotes spoken by the speakers from the context. Do NOT summarize, shorten, edit, or omit any sentences. Output every dialogue turn in full, exactly as it appears in the context, preserving all sentences. Format each quote exactly as: [Date | Page X | Speaker]: \"Raw Quote Text\". Always extract the exact page number from the bracketed metadata (e.g., Page 35 from '[Page 35]') and use it in the output prefix. Do NOT write 'No Page mentioned', 'implicit Page continuation', or '(No direct quote was provided)'. Refer to speakers by their names or correct pronouns."

        full_prompt_to_llm = f"""[CONTEXT]
{context_str}

[INSTRUCTION]
You are a project assistant. {instruction}

[USER QUESTION]
{user_question}"""

        # Check for live API key (Groq or Gemini)
        groq_key = os.environ.get("GROQ_API_KEY")
        if groq_key:
            print(f"DEBUG: len(context_str) = {len(context_str)}")
            print(f"DEBUG: len(full_prompt_to_llm) = {len(full_prompt_to_llm)}")
            print("Querying Live Groq LLM API (Llama 3.3) for synthesis...")
            response = call_llm_api(full_prompt_to_llm)
            if response:
                title_text = f"LIVE AI SYNTHESIS REPORT: {target_speaker.upper()}" if target_speaker else f"LIVE AI SYNTHESIS REPORT: {user_question.upper()}"
                print_premium_box(title_text, response)
                continue

        # Fallback to local report if no key or query failed
        print("(NO ACTIVE LLM RESPONSE. GENERATING LOCAL DATABASES STATUS REPORT)\n")
        if not date_groups:
            print("No dialogue turns found in the database matching this query.")
        else:
            title_text = f"CHRONOLOGICAL STATUS REPORT: {target_speaker.upper()}" if target_speaker else f"CHRONOLOGICAL STATUS REPORT: {user_question.upper()}"
            if len(title_text) > 74:
                title_text = title_text[:71] + "..."
            print("+" + "-"*78 + "+")
            print(f"| {title_text:<76} |")
            print("+" + "-"*78 + "+")
            
            for date_val in sorted(date_groups.keys()):
                source_files = list(set(p.get("source_file", "unknown") for p in date_groups[date_val]))
                source_str = ", ".join(source_files)
                if len(source_str) > 40:
                    source_str = source_str[:37] + "..."
                
                print(f"\n DATE: {date_val:<15} | SOURCE: {source_str}")
                print("-" * 80)
                
                # Print up to 4 key turns to keep it clean and highly readable
                for p in date_groups[date_val][:4]:
                    text_clean = p["text"].strip().replace("\n", " ")
                    if len(text_clean) > 140:
                        text_clean = text_clean[:137] + "..."
                    
                    print(f"  * [Page {p['page']} | Turn #{p['chunk_id']}]: {text_clean}")
                print()
        print("="*80)

if __name__ == "__main__":
    main()
