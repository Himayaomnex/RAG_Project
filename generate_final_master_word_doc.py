"""
Generates the Single Master Word Document (.docx) containing:
1. System Architecture Flowchart Diagram (Embedded High-Res Image)
2. MoSCoW Prioritization Framework (Must Have, Could Have, Nice to Have)
3. The First Principle of Our RAG System (Human-Readable Core Truths)
"""
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = docx.Document()

# Page Margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(4)
r_t = p_title.add_run("ENTERPRISE MULTI-AGENT RAG SYSTEM")
r_t.font.name = "Arial"
r_t.font.size = Pt(20)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(14, 56, 122)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(16)
r_s = p_sub.add_run("Master Technical Report: System Architecture, MoSCoW Framework & First Principle")
r_s.font.name = "Arial"
r_s.font.size = Pt(11)
r_s.font.italic = True
r_s.font.color.rgb = RGBColor(90, 105, 120)

# Section 1: System Architecture Diagram Image
h1 = doc.add_paragraph()
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(6)
r1 = h1.add_run("1. System Architecture Flowchart Diagram")
r1.font.name = "Arial"
r1.font.size = Pt(14)
r1.font.bold = True
r1.font.color.rgb = RGBColor(14, 56, 122)

img_path = r"C:\Users\Omnex\.gemini\antigravity-ide\brain\3f8efc03-ec37-4d57-a983-6ad20233a8dc\updated_direct_metadata_architecture_flowchart_1785400357963.png"
if os.path.exists(img_path):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_after = Pt(16)
    p_img.add_run().add_picture(img_path, width=Inches(6.2))

# Section 2: MoSCoW Framework
h2 = doc.add_paragraph()
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(6)
r2 = h2.add_run("2. Project Prioritization Framework (MoSCoW)")
r2.font.name = "Arial"
r2.font.size = Pt(14)
r2.font.bold = True
r2.font.color.rgb = RGBColor(14, 56, 122)

def add_mh_heading(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0, 0, 0)

add_mh_heading("Must Have")
must_bullets = [
    "Training Transcript Ingestion from Word Documents (.docx): Primary input source parsing and indexing of Microsoft Teams meeting transcripts.",
    "User Scoping (Scope Before Search): Scope all queries, vectors, and responses strictly to caller identity (user_id / speaker) before retrieval begins.",
    "Semantic Transcript Retrieval using Qdrant: 2,843 384-dimensional vector chunks indexed with Cosine similarity search.",
    "Persistent SHA-256 Embedding Cache (emb_cache): Sub-4ms instant vector lookups at $0.00 cost (Implemented).",
    "Grounded AI Responses with Proof: Grounded answers generated strictly from retrieved transcript context.",
    "Exact Source Citations: Citing Date, Page, and Speaker [Date | Page | Speaker] with zero guessing.",
    "FastAPI REST APIs & Agent Access: Programmatic endpoints for multi-agent dispatching."
]
for b in must_bullets:
    add_bullet(b)

add_mh_heading("Should Have")
should_bullets = [
    "GitHub Repository Integration & Automatic Code Review: Connecting directly to GitHub repositories for automatic code reviews so trainees focus on building rather than debugging.",
    "Interactive Code Debugging & Guided Learning Assistant: AI guidance explaining how trainees can debug code to foster learning through problem-solving.",
    "Automatic Folder Watcher (auto_folder_watcher.py): Background daemon automatically indexing new transcript files live.",
    "Quick Preset Action Buttons: Interactive buttons for Quiz Generation, Reading Topics, Action Items, and Spoken Quotes."
]
for b in should_bullets:
    add_bullet(b)

add_mh_heading("Could Have (Future Productivity Enhancements)")
could_bullets = [
    "Microsoft Teams Live Transcript Ingestion via Graph API.",
    "Whisper-Based Speech-to-Text running locally on laptops.",
    "Live In-Meeting Real-Time Prompt Assistant (Teams / Slack).",
    "Export Evaluation Reports directly to Word/PDF.",
    "Visual Analytics Dashboard for team participation metrics.",
    "Qdrant Cloud Synchronization for remote team access."
]
for b in could_bullets:
    add_bullet(b)

# Section 3: First Principle
h3 = doc.add_paragraph()
h3.paragraph_format.space_before = Pt(14)
h3.paragraph_format.space_after = Pt(6)
r3 = h3.add_run("3. The First Principle of Our RAG System")
r3.font.name = "Arial"
r3.font.size = Pt(14)
r3.font.bold = True
r3.font.color.rgb = RGBColor(14, 56, 122)

fp_bullets = [
    ("Understand Meaning, Not Just Keywords: ", "The system uses 384-dimensional semantic vector search to understand the intent behind a query, even when users phrase it differently."),
    ("Scope Before Search: ", "Every query is first scoped to the appropriate user, speaker, or meeting context before any retrieval begins to ensure privacy and relevance."),
    ("Direct Metadata Payload Filtering (Speaker & Date Payloads): ", "Rather than running pre-processing data transformation pipelines, the system leverages native Qdrant metadata payload filtering ('speaker' and 'date' fields) combined with semantic vector search to resolve names and dates naturally during retrieval."),
    ("No Guessing, Only Evidence: ", "Every response is grounded in retrieved transcript content with citations [Date | Page | Speaker]. If no evidence exists, the system returns 'Information Not Available'.")
]

for title, desc in fp_bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r_t = p.add_run(title)
    r_t.font.name = "Calibri"
    r_t.font.bold = True
    r_t.font.size = Pt(11)
    r_d = p.add_run(desc)
    r_d.font.name = "Calibri"
    r_d.font.size = Pt(11)

out_file = "SIDDHARTH_MOSCOW_MASTER_REPORT.docx"
try:
    doc.save(out_file)
    print(f"Updated Master Word Document: {out_file}")
except Exception:
    out_file = "ENTERPRISE_RAG_MASTER_REPORT.docx"
    doc.save(out_file)
    print(f"Updated Master Word Document: {out_file}")

# Auto-copy to artifacts
import shutil
artifact_dest = r"C:\Users\Omnex\.gemini\antigravity-ide\brain\3f8efc03-ec37-4d57-a983-6ad20233a8dc\SIDDHARTH_MOSCOW_MASTER_REPORT.docx"
try:
    shutil.copy(out_file, artifact_dest)
except Exception:
    pass
